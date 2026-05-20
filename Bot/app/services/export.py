import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings
from app.core.logging_config import logger
from app.utils.time_utils import format_seconds_to_hms, format_seconds_to_srt

def generate_plain_txt(
    segments: List[Dict[str, Any]], 
    output_path: Path
) -> Path:
    """
    Generates plain transcript TXT file with timestamps.
    Format: [00:00:01 - 00:00:06] Text here
    """
    logger.info(f"Generating plain text export at {output_path}...")
    with open(output_path, "w", encoding="utf-8") as f:
        for seg in segments:
            start_hms = format_seconds_to_hms(seg["start_time"])
            end_hms = format_seconds_to_hms(seg["end_time"])
            f.write(f"[{start_hms} - {end_hms}] {seg['text']}\n")
    return output_path

def generate_speaker_txt(
    segments: List[Dict[str, Any]], 
    speaker_map: Dict[str, str], 
    output_path: Path
) -> Path:
    """
    Generates transcript TXT file with timestamps and speaker labels.
    Format: [00:00:01 - 00:00:06] Speaker Name: Text here
    """
    logger.info(f"Generating speaker text export at {output_path}...")
    with open(output_path, "w", encoding="utf-8") as f:
        for seg in segments:
            start_hms = format_seconds_to_hms(seg["start_time"])
            end_hms = format_seconds_to_hms(seg["end_time"])
            raw_speaker = seg.get("speaker_label", "Speaker 1")
            # Use mapped name if available, otherwise fallback to raw label
            speaker_name = speaker_map.get(raw_speaker, raw_speaker)
            f.write(f"[{start_hms} - {end_hms}] {speaker_name}: {seg['text']}\n")
    return output_path

def generate_srt(
    segments: List[Dict[str, Any]], 
    output_path: Path
) -> Path:
    """
    Generates standard SRT subtitles file.
    """
    logger.info(f"Generating SRT subtitles export at {output_path}...")
    with open(output_path, "w", encoding="utf-8") as f:
        for idx, seg in enumerate(segments, start=1):
            start_srt = format_seconds_to_srt(seg["start_time"])
            end_srt = format_seconds_to_srt(seg["end_time"])
            text = seg["text"]
            
            f.write(f"{idx}\n")
            f.write(f"{start_srt} --> {end_srt}\n")
            f.write(f"{text}\n\n")
    return output_path

def generate_docx(
    job_metadata: Dict[str, Any],
    segments: List[Dict[str, Any]],
    speaker_map: Dict[str, str],
    summary: Optional[str],
    output_path: Path
) -> Path:
    """
    Generates a professional styled Word DOCX document.
    """
    logger.info(f"Generating professional DOCX report at {output_path}...")
    doc = Document()

    # Document Title Styling
    title = doc.add_paragraph()
    title_run = title.add_run("Transcription & Analysis Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # Deep Slate Blue
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    
    metadata = [
        ("Original File Name", str(job_metadata.get("original_filename", "Unknown"))),
        ("Duration", format_seconds_to_hms(job_metadata.get("duration_seconds", 0.0))),
        ("Detected Language", str(job_metadata.get("language_code", "Unknown")).upper()),
        ("Speakers Detected", str(job_metadata.get("speakers_count", 1))),
        ("Processed Date", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
    ]

    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)

    doc.add_paragraph()  # spacing

    # Summary Section (if available)
    if summary:
        h_summary = doc.add_paragraph()
        h_summary_run = h_summary.add_run("Executive Summary & Action Items")
        h_summary_run.font.name = 'Arial'
        h_summary_run.font.size = Pt(16)
        h_summary_run.font.bold = True
        h_summary_run.font.color.rgb = RGBColor(31, 78, 121)

        # Parse markdown summary roughly (split by lines)
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(4)
            
            # Simple bold parsing (e.g. **text**)
            parts = line.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 != 0:
                    run.bold = True
                
                # Check for bullet points
                if part.strip().startswith("-") or part.strip().startswith("*"):
                    p.paragraph_format.left_indent = Pt(18)

        doc.add_paragraph()

    # Transcript Section
    h_trans = doc.add_paragraph()
    h_trans_run = h_trans.add_run("Transcript")
    h_trans_run.font.name = 'Arial'
    h_trans_run.font.size = Pt(16)
    h_trans_run.font.bold = True
    h_trans_run.font.color.rgb = RGBColor(31, 78, 121)

    for seg in segments:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        # Time stamp run
        start_hms = format_seconds_to_hms(seg["start_time"])
        end_hms = format_seconds_to_hms(seg["end_time"])
        time_run = p.add_run(f"[{start_hms} - {end_hms}] ")
        time_run.font.italic = True
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Speaker run
        raw_speaker = seg.get("speaker_label", "Speaker 1")
        speaker_name = speaker_map.get(raw_speaker, raw_speaker)
        spk_run = p.add_run(f"{speaker_name}: ")
        spk_run.bold = True
        spk_run.font.color.rgb = RGBColor(46, 117, 182)
        
        # Text run
        p.add_run(seg["text"])

    doc.save(str(output_path))
    return output_path

def generate_technical_json(
    job_metadata: Dict[str, Any],
    segments: List[Dict[str, Any]],
    output_path: Path
) -> Path:
    """
    Generates a full technical metadata JSON report.
    """
    logger.info(f"Generating JSON metadata export at {output_path}...")
    
    report = {
        "job_id": job_metadata.get("id"),
        "original_filename": job_metadata.get("original_filename"),
        "duration_seconds": job_metadata.get("duration_seconds"),
        "whisper_model": job_metadata.get("whisper_model"),
        "language_detected": job_metadata.get("language_code"),
        "speakers_detected_count": job_metadata.get("speakers_count"),
        "processing_time_seconds": job_metadata.get("processing_time_seconds"),
        "status": job_metadata.get("status"),
        "error_message": job_metadata.get("error_message"),
        "processed_at": datetime.utcnow().isoformat(),
        "segments": segments
    }
    
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)
        
    return output_path
